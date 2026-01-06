import streamlit as st 
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
from difflib import get_close_matches
import os

st.set_page_config(page_title="Victura Technologies - SAP GRC", layout="wide", page_icon="🔐")

st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
*{font-family:'Inter',sans-serif!important}
.main{background:linear-gradient(135deg,#0f172a 0%,#1e3a8a 50%,#4c1d95 100%);padding:2rem}
.stApp{background:transparent}
div[data-testid="stFileUploader"]{background:white;padding:1.8rem;border-radius:15px;box-shadow:0 10px 30px rgba(0,0,0,0.2);border:3px solid #1e3a8a;transition:all 0.3s}
div[data-testid="stFileUploader"]:hover{border-color:#4c1d95;box-shadow:0 15px 40px rgba(76,29,149,0.3);transform:translateY(-2px)}
div[data-testid="stFileUploader"] label,
div[data-testid="stFileUploader"] span,
div[data-testid="stFileUploader"] p,
section[data-testid="stFileUploadDropzone"],
section[data-testid="stFileUploadDropzone"] span{color:#000000!important;font-weight:600!important;font-size:0.85rem!important}
section[data-testid="stFileUploadDropzone"] small,
[data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] p{color:#000000!important;font-weight:600!important;font-size:0.8rem!important}
div[data-testid="stFileUploader"] small{color:#1e3a8a!important;font-weight:500!important;font-size:0.75rem!important}
section[data-testid="stSidebar"] div[data-testid="stFileUploader"] *{color:#000000!important;font-size:0.85rem!important}
section[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] *{color:#000000!important;font-size:0.85rem!important}
.uploadedFileName{color:#000000!important;font-weight:600!important;font-size:0.85rem!important}
.stTabs [data-baseweb="tab-list"]{gap:1.5rem;background:white;padding:1.8rem;border-radius:15px;box-shadow:0 8px 25px rgba(0,0,0,0.15)}
.stTabs [data-baseweb="tab"]{height:60px;background:linear-gradient(135deg,#1e3a8a,#3b82f6);color:white;border-radius:12px;padding:0 3rem;font-weight:700;font-size:17px;transition:all 0.3s;text-transform:uppercase;letter-spacing:1px}
.stTabs [data-baseweb="tab"]:hover{background:linear-gradient(135deg,#4c1d95,#7c3aed);transform:translateY(-3px);box-shadow:0 8px 20px rgba(76,29,149,0.4)}
.stTabs [aria-selected="true"]{background:linear-gradient(135deg,#4c1d95,#7c3aed);box-shadow:0 8px 25px rgba(76,29,149,0.5)}
h1,h2,h3{color:white!important;text-shadow:3px 3px 10px rgba(0,0,0,0.4);font-weight:800}
section[data-testid="stSidebar"] h2,section[data-testid="stSidebar"] h3{color:white!important}
.stTabs [data-testid="stVerticalBlock"] h2{color:#0f172a!important;text-shadow:none!important;background:linear-gradient(135deg,#f8fafc,#e2e8f0);padding:1.5rem;border-radius:12px;margin-bottom:1.5rem;border-left:6px solid #1e3a8a}
.main h1{color:white!important}
/* Fix subheader color in tabs */
.stTabs [data-testid="stVerticalBlock"] h3{color:#0f172a!important;text-shadow:none!important;font-weight:700;padding:0.8rem 0;margin-top:1.5rem}
.stButton button{background:linear-gradient(135deg,#4c1d95,#7c3aed);color:white;font-weight:700;padding:1rem 3rem;border-radius:12px;box-shadow:0 8px 20px rgba(76,29,149,0.4);font-size:17px;text-transform:uppercase;letter-spacing:1.5px;transition:all 0.3s}
.stButton button:hover{background:linear-gradient(135deg,#6d28d9,#8b5cf6);transform:translateY(-4px);box-shadow:0 12px 30px rgba(76,29,149,0.5)}
.victura-header{background:linear-gradient(135deg,#fff,#f1f5f9);padding:3rem;border-radius:25px;box-shadow:0 15px 40px rgba(0,0,0,0.2);margin-bottom:2.5rem;border:4px solid #1e3a8a;position:relative;overflow:hidden}
.victura-header::before{content:'';position:absolute;top:0;left:0;right:0;height:8px;background:linear-gradient(90deg,#0f172a,#1e3a8a 33%,#4c1d95 66%,#7c3aed)}
.victura-logo-text{font-size:3.5rem;font-weight:900;background:linear-gradient(135deg,#0f172a,#1e3a8a 50%,#4c1d95);-webkit-background-clip:text;-webkit-text-fill-color:transparent;margin:0;letter-spacing:-2px;text-shadow:none}
.victura-tagline{color:#1e3a8a;font-size:1.4rem;font-weight:700;margin-top:.8rem;text-shadow:none}
.victura-subtitle{color:#64748b;font-size:1.1rem;margin-top:.5rem;font-weight:500}
.stMetric{background:white;padding:1.8rem;border-radius:15px;box-shadow:0 6px 20px rgba(0,0,0,0.12);border-left:6px solid #1e3a8a}
section[data-testid="stSidebar"]{background:linear-gradient(180deg,#0f172a,#1e3a8a)}
section[data-testid="stSidebar"] h2{color:white!important;font-weight:800}
section[data-testid="stSidebar"] .stImage{background:white;padding:1rem;border-radius:12px;box-shadow:0 4px 15px rgba(255,255,255,0.1)}
section[data-testid="stSidebar"] .stSuccess{background:rgba(34,197,94,0.15);border-left:4px solid #22c55e}
section[data-testid="stSidebar"] .stWarning{background:rgba(251,191,36,0.15);border-left:4px solid #fbbf24}
section[data-testid="stSidebar"] .stInfo{background:rgba(59,130,246,0.15);border-left:4px solid #3b82f6}
section[data-testid="stSidebar"] p{color:#e2e8f0!important}
.stTextInput input{border-radius:10px;border:2.5px solid #cbd5e1;padding:.9rem;font-weight:500;transition:all 0.3s}
.stTextInput input:focus{border-color:#1e3a8a;box-shadow:0 0 0 4px rgba(30,58,138,0.15)}
.victura-footer{background:white;padding:3rem;border-radius:20px;box-shadow:0 10px 30px rgba(0,0,0,0.15);margin-top:3rem;border-top:6px solid #1e3a8a}
.victura-footer h3{color:#0f172a!important;text-shadow:none;font-weight:800}
.victura-footer ul,.victura-footer ol{color:#334155;line-height:1.8}
.victura-footer strong{color:#1e3a8a;font-weight:700}

/* FIXED: Expander sections with proper spacing and font sizes */
div[data-testid="stExpander"]{
    background:white;
    border-radius:10px;
    margin-bottom:1.2rem!important;
    box-shadow:0 2px 8px rgba(0,0,0,0.1);
}
div[data-testid="stExpander"] summary{
    font-size:0.9rem!important;
    font-weight:600!important;
    line-height:1.5!important;
    padding:1rem 1.2rem!important;
    min-height:50px!important;
    display:flex!important;
    align-items:center!important;
    position:relative!important;
    color:#1e3a8a!important;
}
div[data-testid="stExpander"] summary p{
    font-size:0.9rem!important;
    font-weight:600!important;
    margin:0!important;
    padding:0!important;
    line-height:1.5!important;
    color:#1e3a8a!important;
}
div[data-testid="stExpander"] details summary{
    padding:1rem 1.2rem!important;
}
div[data-testid="stExpander"] summary svg{
    margin-right:0.5rem!important;
}
/* Hide the keyboard_arrow_right text that appears */
div[data-testid="stExpander"] summary span[data-testid="stMarkdownContainer"]{
    display:inline-block!important;
    vertical-align:middle!important;
}
div[data-testid="stExpander"] summary::before{
    content:''!important;
}
/* Target and hide the arrow text specifically */
div[data-testid="stExpander"] details > summary > div:first-child{
    display:none!important;
}
div[data-testid="stExpander"] details summary div[class*="streamlitMarkdown"]{
    width:100%!important;
}
div[data-testid="stExpander"] summary > div[data-testid="stMarkdownContainer"]{
    flex:1!important;
    overflow:hidden!important;
}
/* Force hide any text nodes containing keyboard_arrow */
div[data-testid="stExpander"] summary *:not(p):not(strong):not(em){
    font-size:0!important;
}
div[data-testid="stExpander"] summary p,
div[data-testid="stExpander"] summary strong,
div[data-testid="stExpander"] summary em{
    font-size:0.9rem!important;
}
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"]{
    padding:1.2rem!important;
    line-height:1.8!important;
}
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"] p{
    font-size:0.88rem!important;
    line-height:1.8!important;
    margin-bottom:0.8rem!important;
    padding:0.2rem 0!important;
}
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"] strong{
    font-size:0.88rem!important;
    font-weight:600!important;
    color:#1e3a8a!important;
    display:inline-block!important;
    margin-right:0.3rem!important;
}
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"] ul,
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"] ol{
    margin:0.8rem 0!important;
    padding-left:1.8rem!important;
}
div[data-testid="stExpander"] [data-testid="stMarkdownContainer"] li{
    font-size:0.86rem!important;
    line-height:1.8!important;
    margin-bottom:0.6rem!important;
    padding:0.3rem 0!important;
}

/* Success/Info/Warning boxes inside expanders */
div[data-testid="stExpander"] .stSuccess,
div[data-testid="stExpander"] .stInfo,
div[data-testid="stExpander"] .stWarning{
    padding:0.8rem 1.2rem!important;
    margin:0.8rem 0!important;
    border-radius:8px!important;
    font-size:0.86rem!important;
    line-height:1.7!important;
}
div[data-testid="stExpander"] .stSuccess p,
div[data-testid="stExpander"] .stInfo p,
div[data-testid="stExpander"] .stWarning p{
    font-size:0.86rem!important;
    line-height:1.7!important;
    margin:0.4rem 0!important;
}

/* Column headers in expanders */
div[data-testid="stExpander"] .element-container{
    margin-bottom:1rem!important;
}

/* Dataframe styling */
.stDataFrame{
    margin:1rem 0!important;
}

/* Info/Success/Warning messages in main content */
.stAlert{
    padding:1rem!important;
    margin:1rem 0!important;
    border-radius:10px!important;
    font-size:0.92rem!important;
    line-height:1.6!important;
}
.stAlert p{
    font-size:0.92rem!important;
    line-height:1.6!important;
    margin:0.3rem 0!important;
}
</style>""", unsafe_allow_html=True)

# Constants
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, "victura_logo.png")

# Core standard columns for matching (unchanged)
CORE_STANDARD_COLUMNS = {
    'user_id': ['user id', 'userid', 'user', 'id'],
    'user_name': ['user name', 'username', 'name'],
    'Department-N': ['department', 'dept'],
    'Job Profile/Desig.': ['job profile', 'designation', 'desig'],
    'Business Justification': ['business justification', 'justification', 'biz justification'],
    'Plant Code': ['plant code', 'plantcode'],
    'Plant Name': ['plant name', 'plantname'],
    'Plant': ['plant'],
    'role_1': ['role 1', 'role1'],
    'role_2': ['role 2', 'role2'],
    'function_1': ['function 1', 'function1', 'func1'],
    'function_2': ['function 2', 'function2', 'func2'],
    'tcode1': ['tcode 1', 'tcode1'],
    'tcode2': ['tcode 2', 'tcode2'],
    'user_tcode_f1': ['user tcode f1', 'usertcode1'],
    'user_tcode_f2': ['user tcode f2', 'usertcode2'],
    'Risk_ID': ['risk id', 'riskid'],
    'Risk_Summary': ['risk summary', 'summary'],
    'Mitigation_Control': ['mitigation control', 'mitigation'],
    'Monitoring_Process': ['monitoring process'],
    'Monitoring_Frequency': ['monitoring frequency', 'frequency']
}

# Session state initialization
if 'risk_master' not in st.session_state:
    st.session_state.risk_master = None
if 'risk_columns' not in st.session_state:
    st.session_state.risk_columns = []
if 'user_columns' not in st.session_state:
    st.session_state.user_columns = []
if 'risk_mapping' not in st.session_state:
    st.session_state.risk_mapping = {}
# store manual entries
if 'manual_entries' not in st.session_state:
    st.session_state.manual_entries = []


def fuzzy_match_columns(df_columns, standard_mapping):
    """Match columns with fuzzy logic and prevent duplicates"""
    mapping = {}
    used_df_columns = set()
    df_cols_lower = {col: col.lower().strip() for col in df_columns}
    
    # First pass: exact matches
    for standard_col, variants in standard_mapping.items():
        for df_col, df_col_lower in df_cols_lower.items():
            if df_col in used_df_columns:
                continue
            if df_col_lower in [v.lower() for v in variants]:
                mapping[standard_col] = df_col
                used_df_columns.add(df_col)
                break
    
    # Second pass: fuzzy matches
    for standard_col, variants in standard_mapping.items():
        if standard_col in mapping:
            continue
        for df_col, df_col_lower in df_cols_lower.items():
            if df_col in used_df_columns:
                continue
            matches = get_close_matches(df_col_lower, [v.lower() for v in variants], n=1, cutoff=0.8)
            if matches:
                mapping[standard_col] = df_col
                used_df_columns.add(df_col)
                break
    
    return mapping

@st.cache_data
def load_and_process_risk_master(file_bytes, file_name):
    """Load risk master with performance optimization"""
    df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
    
    # Remove duplicate columns
    df = df.loc[:, ~df.columns.duplicated()]
    df.columns = df.columns.str.strip()
    
    # Map standard columns
    column_mapping = fuzzy_match_columns(df.columns, CORE_STANDARD_COLUMNS)
    
    # Get all columns (mapped + unmapped)
    all_columns = list(df.columns)
    
    # Rename only mapped columns
    reverse_mapping = {v: k for k, v in column_mapping.items()}
    df = df.rename(columns=reverse_mapping)
    
    # Ensure required columns exist
    for col in ['Risk_ID', 'Risk_Summary', 'function_1', 'function_2', 'tcode1', 'tcode2',
                'Mitigation_Control', 'Monitoring_Process', 'Monitoring_Frequency']:
        if col not in df.columns:
            df[col] = ''
    
    # Convert all columns to string type for consistency
    for col in df.columns:
        if col not in ['Tcode_1_Set', 'Tcode_2_Set']:
            df[col] = df[col].astype(str).fillna('')
    
    # Optimize string operations with vectorization
    df['function_1'] = df['function_1'].str.strip().str.upper()
    df['function_2'] = df['function_2'].str.strip().str.upper()
    
    # Pre-process tcode sets for faster matching
    df['Tcode_1_Set'] = df['tcode1'].apply(
        lambda x: set([t.strip().upper() for t in str(x).split(',') if t.strip()])
    )
    df['Tcode_2_Set'] = df['tcode2'].apply(
        lambda x: set([t.strip().upper() for t in str(x).split(',') if t.strip()])
    )
    
    # Create lookup key
    df['lookup_key'] = df['function_1'] + '|||' + df['function_2']
    
    return df, column_mapping, all_columns

@st.cache_data
def process_user_data(file_bytes, file_name):
    """Process user data with duplicate prevention"""
    df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
    
    # Remove duplicate columns immediately
    df = df.loc[:, ~df.columns.duplicated()]
    df.columns = df.columns.str.strip()
    
    # Map columns
    column_mapping = fuzzy_match_columns(df.columns, CORE_STANDARD_COLUMNS)
    
    # Get unmapped columns
    mapped_original_cols = set(column_mapping.values())
    unmapped_cols = [col for col in df.columns if col not in mapped_original_cols]
    
    # Rename mapped columns
    reverse_mapping = {v: k for k, v in column_mapping.items()}
    df = df.rename(columns=reverse_mapping)
    
    # Ensure all columns are strings for Arrow compatibility
    for col in df.columns:
        df[col] = df[col].astype(str).fillna('')
    
    return df, column_mapping, unmapped_cols


def prepare_user_dataframe(df):
    """Prepare user dataframe with necessary columns"""
    df = df.copy()
    
    # Ensure required columns exist
    for col in ['function_1', 'function_2', 'tcode1', 'tcode2', 'user_tcode_f1', 'user_tcode_f2']:
        if col not in df.columns:
            df[col] = ''
        df[col] = df[col].astype(str).str.strip().str.upper().fillna('')

    # Ensure Business Justification exists (keep original case, don't uppercase)
    if 'Business Justification' not in df.columns:
        df['Business Justification'] = ''
    else:
        df['Business Justification'] = df['Business Justification'].astype(str).fillna('')
    # Ensure all other columns are strings
    for col in df.columns:
        if col not in ['function_1', 'function_2', 'tcode1', 'tcode2', 'user_tcode_f1', 'user_tcode_f2']:
            df[col] = df[col].astype(str).fillna('')
    
    df['lookup_key'] = df['function_1'] + '|||' + df['function_2']
    return df

def match_with_risk_master_vectorized(user_df, risk_df):
    """Optimized matching with risk master - high performance for large datasets"""
    # Initialize result columns
    result_columns = ['Mitigation_Control', 'Monitoring_Process', 'Monitoring_Frequency',
                      'Risk_ID', 'Risk_Summary', 'Conflicting_Function', 'Conflicting_Tcode']
    for col in result_columns:
        if col not in user_df.columns:
            user_df[col] = ""
    
    # Ensure all columns are strings for Arrow compatibility
    for col in user_df.columns:
        if col not in ['user_tcode_f1_set', 'user_tcode_f2_set', 'Tcode_1_Set', 'Tcode_2_Set']:
            user_df[col] = user_df[col].astype(str).fillna('')
    
    # Create tcode sets for users (vectorized)
    user_df['user_tcode_f1_set'] = user_df['user_tcode_f1'].apply(
        lambda x: set([s.strip().upper() for s in str(x).split(',') if s.strip()])
    )
    user_df['user_tcode_f2_set'] = user_df['user_tcode_f2'].apply(
        lambda x: set([s.strip().upper() for s in str(x).split(',') if s.strip()])
    )
    
    # Build risk lookup dictionary for O(1) access
    risk_lookup = {}
    for _, r in risk_df.iterrows():
        key = r['lookup_key']
        if key not in risk_lookup:
            risk_lookup[key] = []
        risk_lookup[key].append({
            'Risk_ID': str(r['Risk_ID']),
            'Risk_Summary': str(r['Risk_Summary']),
            'Mitigation_Control': str(r['Mitigation_Control']),
            'Monitoring_Process': str(r['Monitoring_Process']),
            'Monitoring_Frequency': str(r['Monitoring_Frequency']),
            'Business Justification': str(r.get('Business Justification', '')),
            'Tcode_1_Set': r['Tcode_1_Set'],
            'Tcode_2_Set': r['Tcode_2_Set'],
            'function_1': str(r['function_1']),
            'function_2': str(r['function_2'])
        })
    
    # Process matches
    for idx, row in user_df.iterrows():
        lookup_key = str(row.get('lookup_key', ''))
        if not lookup_key or lookup_key not in risk_lookup:
            continue
        
        user_tcode_f1 = row['user_tcode_f1_set']
        user_tcode_f2 = row['user_tcode_f2_set']
        
        for risk in risk_lookup[lookup_key]:
            # Check for conflict: intersection on both sides
            if (user_tcode_f1 & risk['Tcode_1_Set']) and (user_tcode_f2 & risk['Tcode_2_Set']):
                user_df.at[idx, 'Risk_ID'] = risk['Risk_ID']
                user_df.at[idx, 'Risk_Summary'] = risk['Risk_Summary']
                user_df.at[idx, 'Mitigation_Control'] = risk['Mitigation_Control']
                user_df.at[idx, 'Monitoring_Process'] = risk['Monitoring_Process']
                user_df.at[idx, 'Monitoring_Frequency'] = risk['Monitoring_Frequency']
                if not user_df.at[idx, 'Business Justification']:
                    user_df.at[idx, 'Business Justification'] = risk['Business Justification']
                
                all_conf_tcodes = risk['Tcode_1_Set'] | risk['Tcode_2_Set']
                user_df.at[idx, 'Conflicting_Tcode'] = ", ".join(sorted(all_conf_tcodes))
                user_df.at[idx, 'Conflicting_Function'] = f"{risk['function_1']} + {risk['function_2']}"
                break
    
    # Clean up temporary columns and ensure all remaining columns are strings
    result_df = user_df.drop(columns=['lookup_key', 'user_tcode_f1_set', 'user_tcode_f2_set'], errors='ignore')
    
    # Final cleanup: ensure all columns are strings for Arrow
    for col in result_df.columns:
        result_df[col] = result_df[col].astype(str).fillna('')
    
    return result_df


def add_logo_to_header(section, logo_path):
    """Add logo to document header (attempts to add picture; falls back to text)"""
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run()
    try:
        if logo_path and os.path.exists(logo_path):
            run.add_picture(logo_path, width=Inches(1.5))
        else:
            header_para.text = "VICTURA TECHNOLOGIES"
            header_para.runs[0].font.bold = True
            header_para.runs[0].font.size = Pt(18)
            header_para.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    except Exception:
        header_para.text = "VICTURA TECHNOLOGIES"
        header_para.runs[0].font.bold = True
        header_para.runs[0].font.size = Pt(18)
        header_para.runs[0].font.color.rgb = RGBColor(30, 58, 138)


def generate_word_report(df, logo_path=None, max_users_detail = 10000, include_all=False):
    """Generate Word report with options to limit detail for very large datasets.
    - If include_all is False and len(df) > max_users_detail, only first max_users_detail users are detailed and executive summary is added.
    - Adding logo to header ensures it appears on every page.
    """
    doc = Document()
    
    # Setup margins and header
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        if logo_path:
            add_logo_to_header(section, logo_path)
    
    # Title
    title = doc.add_heading('SAP GRC Mitigation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para = doc.add_paragraph("Powered by Victura Technologies")
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para.runs[0].font.size = Pt(14)
    company_para.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    company_para.runs[0].font.bold = True
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total_users = len(df)
    matched = len(df[df['Risk_ID'] != ''])

    # Executive summary
    doc.add_paragraph()
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Total users processed: {total_users:,}")
    doc.add_paragraph(f"Total matched (conflicts found): {matched:,}")
    doc.add_paragraph(f"Match rate: {(matched/total_users*100) if total_users>0 else 0:.1f}%")
    
    # Add small stats table
    stats_table = doc.add_table(rows=3, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    stats = [
        ('Total Users', f"{total_users:,}"),
        ('Matched', f"{matched:,}"),
        ('Unmatched', f"{total_users-matched:,}")
    ]
    for i, (label, value) in enumerate(stats):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()

    # Decide how many details to include
    if not include_all and total_users > max_users_detail:
        doc.add_heading('Detailed sample (first {} users)'.format(max_users_detail), 2)
        detail_df = df.head(max_users_detail)
        doc.add_paragraph('Full dataset is available in the Excel export.')
    else:
        doc.add_heading('Detailed User Information', 2)
        detail_df = df
    
    # Define columns to present succinctly
    present_cols = ['user_id', 'user_name', 'Department-N', 'Job Profile/Desig.','Business Justification', 'role_1', 'role_2',
                    'function_1', 'function_2', 'user_tcode_f1', 'user_tcode_f2',
                    'Risk_ID', 'Risk_Summary', 'Mitigation_Control', 'Monitoring_Process', 'Monitoring_Frequency']
    present_cols = [c for c in present_cols if c in detail_df.columns]
    
    # Add details per user but attempt to be memory efficient
    for idx, row in detail_df.iterrows():
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(f"User: {row.get('user_name', 'N/A')} ({row.get('user_id','')})", 2)
        info = []
        for c in present_cols:
            info.append((c.replace('_', ' ').title(), str(row.get(c, ''))))
        tbl = doc.add_table(rows=len(info), cols=2)
        tbl.style = 'Light Grid Accent 1'
        for i, (lbl, val) in enumerate(info):
            tbl.rows[i].cells[0].text = str(lbl)
            tbl.rows[i].cells[1].text = str(val)
            tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    # Footer note if truncated
    if not include_all and total_users > max_users_detail:
        doc.add_page_break()
        doc.add_paragraph()
        doc.add_heading('Note', 2)
        doc.add_paragraph(f"This Word report includes a detailed sample of the first {max_users_detail:,} users. "
                          "The complete dataset is available in the accompanying Excel export to ensure performance for very large datasets.")
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Header (unchanged)
st.markdown('''<div class="victura-header">
<h1 class="victura-logo-text">VICTURA TECHNOLOGIES</h1>
<p class="victura-tagline">SAP GRC Mitigation Report Generator</p>
<p class="victura-subtitle">Enterprise Compliance & Risk Management Solution</p>
</div>''', unsafe_allow_html=True)

# Sidebar (unchanged functionality, small additions)
with st.sidebar:
    st.header("📁 Risk Master Data")

    # 🖼 STEP 4: SHOW LOGO
    if os.path.exists(LOGO_PATH):
        st.success("✅ Company Logo Loaded")
        st.image(LOGO_PATH, width=150)
    else:
        st.warning("⚠️ Logo not found at configured path")
        st.code(LOGO_PATH, language="text")

    
    risk_file = st.file_uploader("Upload Risk Master Excel", type=['xlsx', 'xls'])
    if risk_file:
        try:
            with st.spinner("Loading Risk Master..."):
                risk_df, risk_mapping, risk_cols = load_and_process_risk_master(
                    risk_file.read(), risk_file.name
                )
                st.session_state.risk_master = risk_df
                st.session_state.risk_mapping = risk_mapping
                st.session_state.risk_columns = risk_cols
            
            st.success(f"✅ Loaded: {len(risk_df):,} records")
            st.info(f"📊 Columns: {len(risk_cols)}")
            
            with st.expander("📊 Column Mapping"):
                mapping_df = pd.DataFrame([
                    {"Standard": k, "Matched": v} 
                    for k, v in risk_mapping.items()
                ])
                st.dataframe(mapping_df, use_container_width=True)
            
            # Show unmapped columns
            mapped_originals = set(risk_mapping.values())
            unmapped = [col for col in risk_cols if col not in mapped_originals]
            if unmapped:
                with st.expander(f"ℹ️ Additional Columns ({len(unmapped)})"):
                    for col in unmapped:
                        st.write(f"• {col}")
        
        except Exception as e:
            st.error(f"❌ Error loading Risk Master: {str(e)}")

# Tabs
tab1, tab2 = st.tabs(["📝 Manual Entry", "📤 Bulk Upload"])

with tab1:
    st.header("Manual User Entry")
    
    if st.session_state.risk_master is None:
        st.warning("⚠️ Please upload Risk Master file first")
    else:
        st.info("Enter user details in the compact form below. You can 'Add' multiple manual entries then generate/download reports for all added entries.")
        # Compact form fields required by user
        c1, c2 = st.columns(2)
        with c1:
            user_id = st.text_input('User ID', key='manual_user_id')
            user_name = st.text_input('User Name', key='manual_user_name')
            department = st.text_input('Department', key='manual_department')
            business_justification = st.text_area('Business Justification', key='manual_biz_just', height=80)
            role_1 = st.text_input('Role 1', key='manual_role_1')
            role_2 = st.text_input('Role 2', key='manual_role_2')
        with c2:
            function_1 = st.text_input('Function 1', key='manual_function_1')
            function_2 = st.text_input('Function 2', key='manual_function_2')
            user_tcode_f1 = st.text_input('User Tcodes F1 (comma separated)', key='manual_user_tcode_f1')
            user_tcode_f2 = st.text_input('User Tcodes F2 (comma separated)', key='manual_user_tcode_f2')
        
        add_col1, add_col2, add_col3 = st.columns([1,1,1])
        with add_col1:
            if st.button('➕ Add Entry'):
                entry = {
                    'user_id': user_id,
                    'user_name': user_name,
                    'Department-N': department,
                    'role_1': role_1,
                    'role_2': role_2,
                    'function_1': function_1,
                    'function_2': function_2,
                    'user_tcode_f1': user_tcode_f1,
                    'user_tcode_f2': user_tcode_f2
                }
                st.session_state.manual_entries.append(entry)
                st.success('Entry added to manual list')
        with add_col2:
            if st.button('🧹 Clear Entries'):
                st.session_state.manual_entries = []
                st.success('Manual entries cleared')
        with add_col3:
            if st.button('🔄 Generate Report for Current Entry'):
                # Single entry immediate processing
                single_df = pd.DataFrame([{
                    'user_id': user_id,
                    'user_name': user_name,
                    'Department-N': department,
                    'Business Justification': business_justification,
                    'role_1': role_1,
                    'role_2': role_2,
                    'function_1': function_1,
                    'function_2': function_2,
                    'user_tcode_f1': user_tcode_f1,
                    'user_tcode_f2': user_tcode_f2
                }])
                single_df = prepare_user_dataframe(single_df)
                processed_df = match_with_risk_master_vectorized(single_df, st.session_state.risk_master)
                st.success('✅ Report generated for current entry')
                st.dataframe(processed_df, use_container_width=True)
                # Downloads
                colx, colw = st.columns(2)
                with colx:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        processed_df.to_excel(writer, index=False)
                    excel_buffer.seek(0)
                    st.download_button('📥 Download Excel (Current Entry)', excel_buffer,
                                       f"SAP_GRC_Manual_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                       'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                with colw:
                    logo_to_use = LOGO_PATH if os.path.exists(LOGO_PATH) else None
                    word_buffer = generate_word_report(processed_df, logo_to_use, max_users_detail=500, include_all=True)
                    st.download_button('📄 Download Word (Current Entry)', word_buffer,
                                       f"SAP_GRC_Manual_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                       'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        st.markdown('---')
        st.subheader('Manual Entries Queue')
        if st.session_state.manual_entries:
            queue_df = pd.DataFrame(st.session_state.manual_entries)
            st.dataframe(queue_df, use_container_width=True)
            q1, q2 = st.columns(2)
            with q1:
                if st.button('🔄 Process All Queue'):
                    proc_df = prepare_user_dataframe(queue_df)
                    processed_df = match_with_risk_master_vectorized(proc_df, st.session_state.risk_master)
                    st.success('✅ Queue processed')
                    st.dataframe(processed_df, use_container_width=True)
                    # provide downloads for full processed queue
                    buf = io.BytesIO()
                    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                        processed_df.to_excel(writer, index=False)
                    buf.seek(0)
                    st.download_button('📥 Download Excel (All Manual Queue)', buf,
                                       f"SAP_GRC_ManualQueue_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                       'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    logo_to_use = LOGO_PATH if os.path.exists(LOGO_PATH) else None
                    word_buf = generate_word_report(processed_df, logo_to_use, max_users_detail=1000, include_all=False)
                    st.download_button('📄 Download Word (Sample Detailed)', word_buf,
                                       f"SAP_GRC_ManualQueue_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                       'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            with q2:
                if st.button('📤 Export Queue to Excel (raw entries)'):
                    raw_buf = io.BytesIO()
                    queue_df.to_excel(raw_buf, index=False)
                    raw_buf.seek(0)
                    st.download_button('📥 Download Raw Queue Excel', raw_buf,
                                       f"SAP_GRC_ManualQueue_Raw_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                       'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        else:
            st.info('No manual entries queued yet. Use the form above and click ➕ Add Entry to queue entries.')

with tab2:
    st.header("Bulk User Upload")
    st.info("📋 Upload any Excel file - columns will be automatically mapped!")
    
    user_file = st.file_uploader("Upload User Conflict Report", type=['xlsx', 'xls'], key='bulk_upload')
    
    if user_file:
        try:
            with st.spinner("Processing file..."):
                user_df, user_mapping, unmapped = process_user_data(
                    user_file.read(), user_file.name
                )
                st.session_state.user_columns = list(user_df.columns)
            
            st.success(f"✅ Loaded: {len(user_df):,} records")
            st.info(f"📊 Columns: {len(user_df.columns)}")
            
            col1, col2 = st.columns(2)
            with col1:
                with st.expander("📊 Column Mapping"):
                    mapping_df = pd.DataFrame([
                        {"Standard": k, "Matched": v}
                        for k, v in user_mapping.items()
                    ])
                    st.dataframe(mapping_df, use_container_width=True)
            
            with col2:
                # Check for critical missing columns
                critical_cols = ['function_1', 'function_2', 'user_tcode_f1', 'user_tcode_f2']
                missing_critical = [col for col in critical_cols if col not in user_mapping.keys()]
                
                if missing_critical:
                    st.error(f"⚠️ Missing Critical Columns for Matching: {len(missing_critical)}")
                    with st.expander("❌ Missing Required Columns", expanded=True):
                        st.warning("**These columns are REQUIRED for risk matching:**")
                        for col in missing_critical:
                            variants = ", ".join(CORE_STANDARD_COLUMNS.get(col, []))
                            st.write(f"• **{col}** (or similar: {variants})")
                        st.info("💡 Your Excel should contain columns with names like these variants")
                
                if unmapped:
                    with st.expander(f"ℹ️ Additional Columns ({len(unmapped)})"):
                        st.write("**These columns don't match standard names but will be included:**")
                        for col in unmapped[:20]:
                            st.write(f"• {col}")
                        if len(unmapped) > 20:
                            st.write(f"... and {len(unmapped)-20} more")
            
            with st.expander("👁️ Data Preview"):
                st.dataframe(user_df.head(10), use_container_width=True)
            
            # Show all columns in the uploaded file
            with st.expander("🔍 All Columns in Your File"):
                st.write("**Original column names from your Excel file:**")
                temp_df = pd.read_excel(io.BytesIO(user_file.getvalue()), nrows=0)
                col_list = temp_df.columns.tolist()
                
                col_display = st.columns(3)
                for idx, col in enumerate(col_list):
                    with col_display[idx % 3]:
                        is_mapped = col in user_mapping.values()
                        if is_mapped:
                            mapped_to = [k for k, v in user_mapping.items() if v == col][0]
                            st.success(f"✅ **{col}**\n→ {mapped_to}")
                        else:
                            st.info(f"ℹ️ **{col}**\n→ Not mapped")
            
            if st.button("🔄 Process Bulk Upload", type="primary", use_container_width=True):
                if st.session_state.risk_master is None:
                    st.error("⚠️ Please upload Risk Master file first!")
                else:
                    with st.spinner("⚙️ Processing bulk upload..."):
                        user_df_prepared = prepare_user_dataframe(user_df)
                        processed_df = match_with_risk_master_vectorized(
                            user_df_prepared, st.session_state.risk_master
                        )
                    
                    st.success("✅ Processing complete!")
                    
                    # Show statistics
                    c1, c2, c3, c4 = st.columns(4)
                    matched = len(processed_df[processed_df['Risk_ID'] != ''])
                    with c1:
                        st.metric("📊 Total Users", f"{len(processed_df):,}")
                    with c2:
                        st.metric("✅ Matched", f"{matched:,}")
                    with c3:
                        st.metric("❌ Unmatched", f"{len(processed_df)-matched:,}")
                    with c4:
                        match_rate = (matched/len(processed_df)*100) if len(processed_df) > 0 else 0
                        st.metric("📈 Match Rate", f"{match_rate:.1f}%")
                    
                    # Display results
                    st.dataframe(processed_df, height=400, use_container_width=True)
                    
                    # Download buttons
                    col1, col2 = st.columns(2)
                    with col1:
                        excel_buffer = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                            processed_df.to_excel(writer, index=False)
                        excel_buffer.seek(0)
                        st.download_button(
                            "📥 Download Excel",
                            excel_buffer,
                            f"SAP_GRC_Bulk_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key='bulk_excel',
                            use_container_width=True
                        )
                    
                    with col2:
                        with st.spinner("📄 Generating Word report..."):
                            logo_to_use = LOGO_PATH if os.path.exists(LOGO_PATH) else None
                            # For very large uploads we generate a sample-detailed Word and full Excel
                            word_buffer = generate_word_report(processed_df, logo_to_use, max_users_detail=1000, include_all=False)
                        st.download_button(
                            "📄 Download Word",
                            word_buffer,
                            f"SAP_GRC_Bulk_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key='bulk_word',
                            use_container_width=True
                        )
        
        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")

# Footer (unchanged - instructions preserved)
st.markdown("""<div class="victura-footer">
<h3>📋 Instructions</h3>
<ol>
<li><strong>Upload Risk Master</strong> - Excel file with risk definitions (columns auto-mapped)</li>
<li><strong>Choose Method</strong>:
   <ul>
   <li><strong>Manual Entry</strong>: Fill compact form (User ID, Name, Functions, Tcodes) and Add to queue or Generate report for a single entry</li>
   <li><strong>Bulk Upload</strong>: Upload user conflict report Excel</li>
   </ul>
</li>
<li><strong>Process & Download</strong> - Generate Excel and Word reports instantly</li>
</ol>

<h3>📊 Required Columns for User Conflict Sheet</h3>
<p><strong>Your Excel MUST contain these columns (or similar names):</strong></p>
<ul>
<li><strong>function_1</strong> (or: Function 1, Func1) - First conflicting function</li>
<li><strong>function_2</strong> (or: Function 2, Func2) - Second conflicting function</li>
<li><strong>user_tcode_f1</strong> (or: User Tcode F1, UserTcode1) - User's tcodes for function 1</li>
<li><strong>user_tcode_f2</strong> (or: User Tcode F2, UserTcode2) - User's tcodes for function 2</li>
<li><strong>user_id, user_name, department</strong> - Basic user info</li>
<li><strong>Business Justification</strong> (or: Justification, Biz Justification) - Business reason for access</li>
<li><strong>role_1, role_2</strong> (or: Role 1, Role 2, Roles) - User roles</li>
</ul>
<p style="color:#dc2626;font-weight:600">⚠️ Without function and tcode columns, matching will NOT work!</p>

<h3>⚡ Key Features</h3>
<ul>
<li><strong>Smart Column Mapping</strong>: Fuzzy matching - no exact column names needed</li>
<li><strong>Manual Entry Queue</strong>: Add multiple manual users, process them together, export raw or processed data</li>
<li><strong>Large Dataset Support</strong>: Optimized matching for 60,000+ rows and Word export that avoids creating an enormous Word file by providing a detailed sample and full Excel export for very large datasets</li>
<li><strong>Dynamic Forms</strong>: Manual entry adapts to required compact fields while keeping interface design</li>
<li><strong>Duplicate Prevention</strong>: Automatic duplicate column removal</li>
<li><strong>High Performance</strong>: Optimized for large datasets</li>
<li><strong>Professional Reports</strong>: Branded Word docs with company logo on every page (if configured)
</ul>
<div style="margin-top:2rem;text-align:center;padding-top:2rem;border-top:2px solid #e2e8f0">
<p style="color:#0f172a;font-weight:800;font-size:1.2rem;margin:0">VICTURA TECHNOLOGIES</p>
<small style="color:#64748b">Enterprise SAP GRC Solutions | Version 2.5 Enhanced</small>
</div>
</div>""", unsafe_allow_html=True)








